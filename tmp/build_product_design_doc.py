from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\alexa\Code\Competitions\LIDM\LIDM-Gizi-Sam")
OUT = ROOT / "output" / "product-design"
SHOTS = OUT / "screenshots"
OUT.mkdir(parents=True, exist_ok=True)

GREEN = "0F5C4D"
DARK = "17352F"
MINT = "CFF7DF"
PALE = "EFF8F3"
GRAY = "5F6B66"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.line_spacing = 1.25
normal.paragraph_format.space_after = Pt(6)

for style_name, size, before, after in [("Heading 1", 14, 16, 8), ("Heading 2", 12, 12, 6), ("Heading 3", 11, 8, 4)]:
    st = styles[style_name]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(DARK)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tcMar.find(qn(f"w:{m}"))
        if el is None:
            el = OxmlElement(f"w:{m}")
            tcMar.append(el)
        el.set(qn("w:w"), str(v)); el.set(qn("w:type"), "dxa")

def set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(9); r.italic = True

def add_picture(filename, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(SHOTS / filename), width=Inches(width))
    add_caption(caption)

def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.26)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    p.add_run(text)

def add_callout(title, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.5)
    c = t.cell(0,0); shade(c, PALE); margins(c, 130, 160, 130, 160)
    c.text = ""
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title); r.bold=True; r.font.color.rgb=RGBColor.from_string(GREEN); r.font.size=Pt(10.5)
    p2 = c.add_paragraph(text); p2.paragraph_format.space_after=Pt(0); p2.paragraph_format.line_spacing=1.16
    for r2 in p2.runs: r2.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(3)
r = p.add_run("4. DESAIN PRODUK")
r.bold = True; r.font.name="Times New Roman"; r.font.size=Pt(16); r.font.color.rgb=RGBColor.from_string(DARK)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(14)
r2 = p2.add_run("NUTRISPHERE - EKOSISTEM PEMBELAJARAN GIZI TERPADU")
r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=RGBColor.from_string(GREEN)

intro = ("NutriSphere dirancang sebagai ekosistem pembelajaran gizi berbasis web yang menyatukan asesmen awal, "
         "sumber ilmiah, simulasi penalaran klinis, pembelajaran berbasis kasus, dan evaluasi kompetensi dalam satu alur. "
         "Desain produk menempatkan mahasiswa sebagai pengguna utama, sedangkan dosen memperoleh ringkasan perkembangan "
         "belajar melalui dashboard. Antarmuka dibuat konsisten, mudah dipelajari, dan dapat diakses melalui peramban tanpa "
         "proses instalasi aplikasi tambahan.")
doc.add_paragraph(intro)
add_picture("01-landing.png", "Gambar 4.1. Tampilan halaman utama NutriSphere pada versi web.")

doc.add_heading("4.1 Desain UI/UX", level=1)
doc.add_paragraph(
    "Desain UI/UX NutriSphere dikembangkan dengan pendekatan user-centered untuk membantu mahasiswa bergerak dari pemetaan "
    "pengetahuan awal menuju latihan penalaran dan evaluasi hasil. Struktur navigasi mempertahankan lima pilar produk pada "
    "sidebar yang sama di setiap halaman, sehingga pengguna tidak perlu mempelajari pola navigasi baru ketika berpindah fitur. "
    "Hierarki visual menonjolkan judul halaman, konteks fitur, tindakan utama, dan status akun secara berurutan."
)

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.columns[0].width = Inches(1.65); table.columns[1].width = Inches(4.85)
for i, h in enumerate(["Aspek Desain", "Implementasi pada NutriSphere"]):
    set_cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF", size=10)
    shade(table.rows[0].cells[i], GREEN)
rows = [
    ("Konsistensi", "Sidebar, kartu, tombol, label fitur, dan status akun memakai pola visual yang seragam pada seluruh halaman."),
    ("Hierarki informasi", "Judul dan konteks halaman ditampilkan terlebih dahulu, diikuti aktivitas utama, informasi pendukung, serta navigasi lanjutan."),
    ("Aksesibilitas", "Kontras hijau gelap dan latar terang, ukuran teks yang terbaca, fokus keyboard, serta struktur tombol yang jelas mendukung kemudahan penggunaan."),
    ("Umpan balik", "Status loading, validasi form, skor, progress kasus, dan pesan kegagalan membantu pengguna memahami hasil setiap tindakan."),
    ("Responsivitas", "Tata letak menyesuaikan desktop dan perangkat bergerak melalui navigasi yang dapat diciutkan serta komponen berbasis kartu."),
    ("Keamanan alur", "Login dan pretest menjadi gerbang akses fitur, sedangkan parameter tujuan membawa pengguna kembali ke halaman yang semula dipilih."),
]
for a,b in rows:
    cells=table.add_row().cells
    set_cell_text(cells[0],a,bold=True,color=DARK); set_cell_text(cells[1],b)
    shade(cells[0],"E7F3EC")
add_caption("Tabel 4.1. Prinsip UI/UX NutriSphere.")

doc.add_heading("4.1.1 Identitas Visual", level=2)
doc.add_paragraph(
    "Identitas visual NutriSphere menggunakan warna hijau tua sebagai warna utama untuk membangun asosiasi dengan kesehatan, "
    "pertumbuhan, dan keberlanjutan. Latar putih-kehijauan menjaga area baca tetap ringan, sedangkan aksen hijau terang dipakai "
    "untuk tombol utama, indikator aktif, dan penanda capaian. Tipografi sans-serif berukuran besar pada judul memperkuat kesan "
    "modern, sementara kartu berujung membulat memberi tampilan ramah dan tidak menyerupai sistem klinis yang kaku."
)
add_bullet("Warna utama: hijau tua untuk navigasi dan identitas produk.")
add_bullet("Warna aksen: hijau terang untuk call-to-action, status aktif, dan indikator keberhasilan.")
add_bullet("Latar: putih dan gradasi hijau sangat muda untuk menjaga keterbacaan konten.")
add_bullet("Ikon: ikon semantik berbasis SVG agar arti fitur dapat dikenali secara cepat dan konsisten.")

doc.add_heading("4.1.2 Arsitektur Informasi dan Alur Pengguna", level=2)
doc.add_paragraph(
    "Arsitektur informasi dibangun berdasarkan lima pilar: NutriSolve, NutriBase, NutriPath, NutriRead, dan NutriQuest. "
    "Alur utama mahasiswa dimulai dari halaman beranda, dilanjutkan login atau registrasi, pretest, eksplorasi materi dan simulasi, "
    "latihan kasus, kemudian posttest. Pretest berfungsi sebagai feature gate global agar aktivitas belajar memiliki data baseline, "
    "sedangkan posttest digunakan sebagai evaluasi akhir tanpa mengunci akses ke fitur lain."
)
flow = doc.add_table(rows=1, cols=4)
flow.alignment=WD_TABLE_ALIGNMENT.CENTER; flow.autofit=False
labels=[("01", "Ukur baseline\nLogin dan pretest"),("02", "Pelajari fondasi\nReferensi dan modul"),("03", "Latih penalaran\nDSS dan kasus 3D"),("04", "Buktikan progres\nPosttest dan dashboard")]
for i,(num,txt) in enumerate(labels):
    c=flow.rows[0].cells[i]; shade(c, MINT if i in (0,3) else PALE); margins(c,120,100,120,100)
    c.text=""; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=p.add_run(num); rr.bold=True; rr.font.color.rgb=RGBColor.from_string(GREEN); rr.font.size=Pt(13)
    pp=c.add_paragraph(txt); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER; pp.paragraph_format.space_after=Pt(0); pp.paragraph_format.line_spacing=1.05
    for rr2 in pp.runs: rr2.font.size=Pt(9); rr2.bold=True
add_caption("Gambar 4.2. Alur pengalaman belajar mahasiswa pada NutriSphere.")
add_callout("Catatan penggunaan", "Hasil antropometri, clinical scanner, dietary mockup, AI Summary, NutriBot, dan AR Patient digunakan sebagai sarana pembelajaran serta decision support edukatif. Fitur tersebut tidak menggantikan asesmen, diagnosis, atau keputusan klinis oleh tenaga profesional.")

doc.add_heading("4.1.3 Pola Interaksi", level=2)
doc.add_paragraph(
    "Interaksi utama menggunakan kartu, tab, slider, form, dan percakapan agar sesuai dengan karakter setiap tugas. Kartu dipakai "
    "untuk memilih pilar atau resource; tab memisahkan kelompok data pasien; form dipakai untuk input asesmen; dan chat digunakan "
    "untuk simulasi umpan balik klinis serta bantuan NutriBot. Seluruh initializer global dirancang tetap aman saat komponen tertentu "
    "tidak tersedia pada suatu halaman."
)
add_picture("02-login.png", "Gambar 4.3. Tampilan login dan registrasi sebagai pintu masuk pengguna.", 5.9)

doc.add_heading("4.2 Penjelasan Fitur", level=1)
doc.add_paragraph(
    "NutriSphere memiliki lima pilar yang saling terhubung. NutriSolve mendukung latihan pengambilan keputusan, NutriBase menyediakan "
    "referensi, NutriPath mengarahkan proses belajar, NutriRead mendukung literasi ilmiah, dan NutriQuest mengukur perkembangan "
    "kompetensi. Penjelasan berikut menggambarkan fungsi yang telah tersedia pada prototipe web."
)

features = [
    ("NutriSolve", "Pusat decision support system untuk menghubungkan data antropometri, tanda klinis, pola makan, dan simulasi pasien 3D. Subfiturnya meliputi Anthropometry Assessment, Clinical Nutrition Screening, Dietary Pattern Assessment, dan AR Patient Visualization."),
    ("NutriBase", "Katalog referensi gizi yang menyediakan pencarian dan filter untuk PDF lokal, regulasi, sumber Drive, serta resource yang masih berstatus placeholder. Pembukaan resource dapat dicatat sebagai event pembelajaran."),
    ("NutriPath", "Hub modul pembelajaran berupa presentasi, video, dan kasus gizi. Klik pada modul mencatat status awal belajar dan persentase progress sebagai indikator aktivitas, bukan bukti bahwa materi telah diselesaikan."),
    ("NutriRead", "Pustaka jurnal dan e-book eksternal yang dilengkapi AI Summary. Teks jurnal dapat diringkas melalui backend Gemini; apabila backend tidak tersedia, sistem memberi respons simulasi berbasis kata kunci dan menandainya sebagai local simulation."),
    ("NutriQuest", "Pusat evaluasi kompetensi yang menghubungkan pretest, posttest, asesmen kasus AR, tracking progress, feature events, dan rekap dashboard dosen."),
]
t = doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
t.columns[0].width=Inches(1.35); t.columns[1].width=Inches(5.15)
for i,h in enumerate(["Pilar", "Fungsi Utama"]): set_cell_text(t.rows[0].cells[i],h,True,"FFFFFF",10); shade(t.rows[0].cells[i],GREEN)
for a,b in features:
    c=t.add_row().cells; set_cell_text(c[0],a,True,DARK); set_cell_text(c[1],b); shade(c[0],"E7F3EC")
add_caption("Tabel 4.2. Ringkasan lima pilar NutriSphere.")

doc.add_heading("4.2.1 NutriSolve - Decision Support System", level=2)
doc.add_paragraph(
    "NutriSolve dirancang sebagai ruang latihan clinical reasoning. Pada Anthropometry Assessment, mahasiswa dapat mempelajari "
    "klasifikasi Z-score untuk balita dan remaja, perhitungan BMI serta rasio lingkar pinggang-panggul pada dewasa, dan estimasi "
    "antropometri pasien rumah sakit. Seluruh perhitungan berjalan di sisi klien sehingga hasil muncul secara langsung setelah input berubah."
)
add_picture("04-antropometri.png", "Gambar 4.4. Antarmuka Anthropometry Assessment pada NutriSolve.")
doc.add_paragraph(
    "Clinical Nutrition Screening memakai kamera untuk mengambil frame dan menampilkan hasil simulasi sesuai skenario anemia atau "
    "stunting. Fitur ini bukan inferensi computer vision. Dietary Pattern Assessment saat ini berupa mockup edukatif, sedangkan AR "
    "Patient Visualization menampilkan model pasien 3D interaktif dan kasus bertingkat dari balita hingga lansia."
)
add_picture("05-ar-patient.png", "Gambar 4.5. Peta progres dan pemilihan level pada AR Patient Visualization.")
add_bullet("Model 3D dapat diputar dan diperbesar melalui kontrol interaktif.")
add_bullet("Kasus menampilkan profil, antropometri, pemeriksaan klinis, dan dietary recall.")
add_bullet("AI Clinical Supervisor memberi umpan balik dialogis terhadap diagnosis dan intervensi yang ditulis mahasiswa.")
add_bullet("Penyelesaian kasus memperbarui streak, level, fase, serta riwayat attempt apabila pengguna terhubung ke Supabase.")

doc.add_heading("4.2.2 NutriBase - Katalog Referensi", level=2)
doc.add_paragraph(
    "NutriBase memusatkan sumber rujukan seperti TKPI, AKG, regulasi BPOM, daftar konversi penyerapan minyak, buku foto makanan, "
    "serta PAGT dan rumus gizi. Pencarian dilakukan berdasarkan judul, sedangkan filter membedakan PDF lokal, tautan Drive, dan "
    "resource placeholder. Desain kartu memuat identitas resource, kategori, dan tindakan buka agar pengguna dapat menemukan "
    "referensi secara cepat."
)

doc.add_heading("4.2.3 NutriPath - Jalur Pembelajaran", level=2)
doc.add_paragraph(
    "NutriPath menyediakan akses terstruktur ke materi presentasi, video, dan studi kasus. Setiap kartu dilengkapi metadata tracking "
    "yang dapat mengubah status learning_progress menjadi in_progress dengan persentase awal. Mekanisme ini membantu dosen melihat "
    "jejak keterlibatan, tetapi tidak mengklaim bahwa materi telah benar-benar ditonton atau diselesaikan."
)

doc.add_heading("4.2.4 NutriRead - Literasi Ilmiah dan AI Summary", level=2)
doc.add_paragraph(
    "NutriRead menjadi pintu masuk ke jurnal, e-book, dan alat AI Summary. Pengguna menempelkan teks jurnal, kemudian sistem melakukan "
    "health check ke backend dan mengirim teks untuk diringkas. Ringkasan diarahkan untuk memuat gambaran eksekutif, temuan penting, "
    "serta evaluasi metodologi dan relevansi dalam Bahasa Indonesia. Jika layanan AI tidak aktif, antarmuka tetap menyediakan fallback "
    "simulasi sehingga alur demonstrasi dapat dilanjutkan dengan label mode yang sesuai."
)
add_picture("06-nutriread-ai.png", "Gambar 4.6. Antarmuka AI Summary pada pilar NutriRead.")

doc.add_heading("4.2.5 NutriQuest - Evaluasi Kompetensi", level=2)
doc.add_paragraph(
    "NutriQuest menghubungkan evaluasi awal dan akhir dengan aktivitas belajar. Pretest dan posttest masing-masing terdiri atas 25 "
    "soal pilihan ganda dari kategori PSG, GDDK, MIPMG, dan MAKRO. Jawaban wajib dipilih sebelum melanjutkan, lalu skor, persentase, "
    "waktu, dan detail jawaban disimpan. Satu pengguna hanya memiliki satu attempt untuk setiap jenis tes."
)
add_picture("07-nutriquest.png", "Gambar 4.7. Halaman pusat evaluasi kompetensi NutriQuest.")
doc.add_paragraph(
    "Dashboard dosen menyajikan identitas mahasiswa, nilai dan waktu pretest-posttest, improvement, jumlah kasus selesai, modul yang "
    "dipelajari, event fitur, dan status belajar. Pada prototipe, akses admin masih memakai password demo di sisi browser dan RPC "
    "dengan secret statis; mekanisme tersebut harus diganti dengan role/claim dan otorisasi server-side sebelum penggunaan produksi."
)

doc.add_heading("4.2.6 Fitur Pendukung", level=2)
support = [
    ("Autentikasi dan profil", "Login email/password, registrasi nama lengkap dan NIM, session Supabase, serta mock user untuk pengembangan lokal."),
    ("Pretest gate", "Mengunci fitur inti sampai pengguna login dan menyelesaikan pretest, sehingga pembelajaran memiliki baseline yang terukur."),
    ("NutriBot", "Widget bantuan global yang menjawab pertanyaan dasar gizi dan navigasi fitur melalui backend chat atau fallback offline."),
    ("Tracking", "Mencatat feature event, pembukaan resource, progress modul, hasil tes, dan attempt kasus secara best-effort."),
    ("Kalkulator makronutrisi", "Slider interaktif pada landing page untuk mengeksplorasi konversi distribusi energi menjadi gram karbohidrat, protein, dan lemak."),
]
t2=doc.add_table(rows=1,cols=2); t2.alignment=WD_TABLE_ALIGNMENT.CENTER; t2.autofit=False
t2.columns[0].width=Inches(1.8); t2.columns[1].width=Inches(4.7)
for i,h in enumerate(["Fitur", "Penjelasan"]): set_cell_text(t2.rows[0].cells[i],h,True,"FFFFFF",10); shade(t2.rows[0].cells[i],GREEN)
for a,b in support:
    c=t2.add_row().cells; set_cell_text(c[0],a,True,DARK); set_cell_text(c[1],b); shade(c[0],"E7F3EC")
add_caption("Tabel 4.3. Fitur pendukung ekosistem NutriSphere.")

add_callout("Kesimpulan desain", "NutriSphere membentuk satu siklus pembelajaran: mengukur baseline, menyediakan sumber belajar, melatih penalaran melalui DSS dan kasus, lalu merekam perkembangan kompetensi. Desain modular memungkinkan setiap pilar dikembangkan secara bertahap tanpa mengubah pola navigasi utama yang telah dipahami pengguna.")

# Header/footer
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
hr=header.add_run("NUTRISPHERE | DESAIN PRODUK"); hr.font.size=Pt(8); hr.font.color.rgb=RGBColor.from_string(GRAY); hr.bold=True
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
fr=footer.add_run("NutriSphere - Ekosistem Pembelajaran Gizi"); fr.font.size=Pt(8); fr.font.color.rgb=RGBColor.from_string(GRAY)

# Keep tables from splitting rows
for table in doc.tables:
    for row in table.rows:
        trPr=row._tr.get_or_add_trPr(); cant=OxmlElement("w:cantSplit"); trPr.append(cant)

target=OUT / "Desain_Produk_NutriSphere.docx"
doc.save(target)
print(target)
